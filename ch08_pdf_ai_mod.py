import streamlit as st
import fitz
from openai import OpenAI
import os
import tempfile
import base64
import shutil
from io import BytesIO
from docx import Document
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.pagesizes import A4
from reportlab.lib.units import inch
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont

# PDF → 이미지 변환 및 base64 인코딩 (메모리 누수 수정)
def convert_pdf_to_base64_images(pdf_data):
    document = None
    temp_dir = None
    try:
        document = fitz.open(stream=pdf_data, filetype="pdf")
        images = []
        base64_images = []
        temp_dir = tempfile.mkdtemp()
        
        for page_num in range(len(document)):
            page = document[page_num]
            # DPI를 높여서 더 선명한 이미지 생성
            pix = page.get_pixmap(dpi=200)
            
            # 이미지 파일로 저장 (미리보기용)
            img_path = os.path.join(temp_dir, f"page_{page_num+1}.png")
            pix.save(img_path)
            images.append(img_path)
            
            # base64 인코딩 (API 전송용)
            img_bytes = pix.tobytes("png")
            base64_img = base64.b64encode(img_bytes).decode('utf-8')
            base64_images.append(base64_img)
            
            # pixmap 메모리 해제
            pix = None
        
        return images, base64_images
        
    except Exception as e:
        st.error(f"PDF 변환 중 오류 발생: {e}")
        return [], []
    finally:
        # 리소스 정리
        if document:
            document.close()
        # 임시 디렉토리는 나중에 정리하기 위해 반환된 images 경로들과 함께 관리

def cleanup_temp_files(temp_paths):
    """임시 파일들을 안전하게 정리"""
    for path in temp_paths:
        try:
            if os.path.exists(path):
                if os.path.isfile(path):
                    os.unlink(path)
                elif os.path.isdir(path):
                    shutil.rmtree(path)
        except Exception as e:
            st.warning(f"임시 파일 정리 중 오류: {e}")

def validate_openai_api_key(api_key):
    """OpenAI API 키 유효성 검사"""
    try:
        client = OpenAI(api_key=api_key)
        # 간단한 API 호출로 키 유효성 검사
        response = client.models.list()
        
        # 사용 가능한 모델 확인 (GPT-5 모델 사용 가능 여부 체크)
        available_models = [model.id for model in response.data]
        
        # GPT-5 모델 사용 가능 여부 확인
        gpt5_available = any("gpt-5" in model for model in available_models)
        
        return client, True, None
    except Exception as e:
        error_msg = str(e)
        if "Incorrect API key" in error_msg:
            return None, False, "잘못된 API 키입니다."
        elif "You exceeded your current quota" in error_msg:
            return None, False, "API 사용 한도를 초과했습니다."
        elif "does not exist" in error_msg:
            return None, False, "선택한 모델에 접근할 수 없습니다. API 키 권한을 확인해주세요."
        else:
            return None, False, f"API 키 검증 오류: {error_msg}"

def extract_context_for_next_page(content, max_length=800):
    """다음 페이지를 위한 문맥 추출"""
    if not content:
        return ""
    
    lines = content.split('\n')
    context_lines = []
    current_length = 0
    
    for line in reversed(lines[-10:]):
        line = line.strip()
        if line and current_length + len(line) < max_length:
            context_lines.insert(0, line)
            current_length += len(line) + 1
        elif current_length > 0:
            break
    
    return '\n'.join(context_lines).strip()

def extract_overlap_context(content, max_length=400):
    """페이지 겹침을 위한 문맥"""
    if not content:
        return "", ""
    
    lines = content.split('\n')
    non_empty_lines = [line.strip() for line in lines if line.strip()]
    
    if not non_empty_lines:
        return "", ""
    
    # 시작 부분
    start_context = []
    start_length = 0
    for line in non_empty_lines[:5]:
        if start_length + len(line) < max_length:
            start_context.append(line)
            start_length += len(line)
        else:
            break
    
    # 끝 부분
    end_context = []
    end_length = 0
    for line in reversed(non_empty_lines[-5:]):
        if end_length + len(line) < max_length:
            end_context.insert(0, line)
            end_length += len(line)
        else:
            break
    
    return '\n'.join(start_context), '\n'.join(end_context)

# GPT Vision API로 이미지 분석 (개별 처리 버전) - 에러 처리 개선
def analyze_single_image_with_context(client, base64_img, prompt_type, model, max_tokens, image_detail, page_num, total_pages, previous_context="", next_page_start=""):
    if not client or not base64_img:
        return None
        
    if prompt_type == "summary":
        context_info = ""
        if previous_context:
            context_info += f"\n**이전 페이지 마지막 내용**:\n{previous_context}\n"
        if next_page_start:
            context_info += f"\n**다음 페이지 시작 내용** (참고용):\n{next_page_start}\n"
            
        system_prompt = f"""당신은 전문 문서 요약 전문가입니다. 현재 {total_pages}페이지 중 {page_num}페이지를 분석하고 있습니다.

📋 **요약 규칙:**
- 개조식으로 요약 (~음, ~했음 어조 사용)
- 핵심 내용과 주요 포인트 중심으로 정리
- 마크다운 형식으로 구조화
- 표나 그래프가 있다면 주요 데이터와 수치를 포함
- 문서의 전체적인 흐름과 맥락 고려
- 중요한 결론이나 시사점 강조
- 페이지 상단/하단의 머리말, 꼬리말, 반복 제목 등은 제외
- 수학 공식이나 수식이 포함된 경우 LaTeX 형식으로 표현하세요
- 본문에서 수식 변수나 기호를 언급할 때도 LaTeX로 표현하세요 (예: "변수 $x$", "$\\beta$ 계수")
- 그림이나 차트가 있는 경우 주요 내용과 데이터를 텍스트로 설명하세요

{context_info}

**주의사항**: 
- 이전/다음 페이지와 연결되는 내용이 있다면 자연스럽게 연결하여 요약하세요
- 페이지 중간에서 끊어진 문장이나 개념이 있다면 완전한 의미로 요약하세요
- 중요한 내용이 누락되지 않도록 충분히 상세하게 요약하세요
- 머리말/꼬리말에 나타나는 반복적인 제목이나 페이지 정보는 제외하세요"""

    else:  # translation
        context_info = ""
        if previous_context:
            context_info += f"\n**이전 페이지 마지막 내용**:\n{previous_context}\n"
        if next_page_start:
            context_info += f"\n**다음 페이지 시작 내용** (참고용):\n{next_page_start}\n"
            
        system_prompt = f"""당신은 고급 전문 번역가입니다. 현재 {total_pages}페이지 중 {page_num}페이지를 번역하고 있습니다.

🌐 **번역 규칙:**
1. **정확성**: 원문의 의미와 뉘앙스를 정확히 보존
2. **전문성**: 학술적/기술적 용어는 적절한 한국어 전문 용어 사용
3. **자연스러움**: 한국어 문체와 어순에 맞게 자연스럽게 번역
4. **구조 보존**: 원문의 문단 구조와 강조점 유지
5. **표와 데이터**: 표는 마크다운 표 형식으로 번역, 수치와 데이터는 정확히 보존
6. **그림과 다이어그램**: 그림 내 텍스트는 모두 번역, 그림 자체는 [그림: 간단한 설명]으로 표현
7. **제목과 소제목**: 적절한 한국어 형식으로 번역
8. **참고문헌**: References, Bibliography 등은 원문 그대로 유지 (섹션 제목만 번역 가능)
9. **머리말/꼬리말**: 페이지 상단/하단의 반복 내용은 제외
10. **수식**: LaTeX 형식으로 표현 (예: $E=mc^2$)
11. **수식 기호 설명**: 본문에서 수식의 변수나 기호를 언급할 때도 LaTeX 사용 (예: "변수 $x$는", "$\\alpha$ 계수", "$\\sigma$ 값")
12. **출력**: 번역문만 제공, 원문과 번역문 병기 금지

{context_info}

**중요**: 
- 이전/다음 페이지와 연결되는 문장이나 문단이 있다면 자연스럽게 연결하여 번역하세요
- 페이지 중간에서 끊어진 문장이 있다면 완전한 문장으로 번역하세요
- 모든 내용을 빠뜨리지 말고 완전히 번역하세요
- 참고문헌 목록이 포함된 경우 해당 부분은 원문 그대로 유지하세요
- 페이지 상단/하단의 머리말, 꼬리말, 반복되는 제목 등은 제외하고 본문만 번역하세요
- 원문을 번역문과 함께 제시하지 말고, 번역문만 제공하세요
- 수학 공식이나 수식은 LaTeX 형식으로 표현하세요
- 본문에서 수식의 변수나 기호를 설명할 때도 LaTeX로 표현하세요 (예: "변수 $x$는...", "$\\alpha$값이...", "$F(x)$ 함수는...")
- 그림이나 다이어그램의 텍스트만 번역하고, 그림 설명은 간단히 추가하세요"""

    messages = [
        {
            "role": "user",
            "content": [
                {
                    "type": "text",
                    "text": f"{system_prompt}\n\n위 규칙에 따라 현재 페이지의 이미지 내용을 정확하고 완전하게 분석해주세요. 모든 텍스트, 표, 그래프를 빠뜨리지 말고 처리하되, 페이지 상단/하단의 머리말이나 꼬리말은 제외하고 본문 내용만 처리하세요. 수식과 본문의 수학 기호 모두 LaTeX 형식으로 표현하세요."
                },
                {
                    "type": "image_url",
                    "image_url": {
                        "url": f"data:image/png;base64,{base64_img}",
                        "detail": image_detail
                    }
                }
            ]
        }
    ]

    try:
        # GPT-5 모델 사용 시 더 높은 타임아웃 설정
        timeout_seconds = 180 if model.startswith("gpt-5") else 60
        
        response = client.chat.completions.create(
            model=model,
            messages=messages,
            max_completion_tokens=max_tokens,
            timeout=timeout_seconds
        )
        
        if response.choices and len(response.choices) > 0:
            choice = response.choices[0]
            if hasattr(choice, 'message') and choice.message and choice.message.content:
                content = choice.message.content.strip()
                if choice.finish_reason == 'length':
                    st.warning(f"⚠️ 페이지 {page_num}: 응답이 토큰 제한으로 잘렸습니다.")
                return content
    except Exception as e:
        error_msg = str(e)
        if "timeout" in error_msg.lower():
            st.error(f"페이지 {page_num} 처리 시간 초과: 네트워크 연결을 확인해주세요.")
        elif "rate limit" in error_msg.lower():
            st.error(f"페이지 {page_num} API 사용량 한도 초과: 잠시 후 다시 시도해주세요.")
        else:
            st.error(f"페이지 {page_num} 처리 오류: {e}")
    
    return None

# 기존 analyze_single_image 함수 (호환성 유지)
def analyze_single_image(client, base64_img, prompt_type, model, max_tokens, image_detail):
    return analyze_single_image_with_context(client, base64_img, prompt_type, model, max_tokens, image_detail, 1, 1, "")

# GPT Vision API로 이미지 분석 - 에러 처리 및 검증 개선
def analyze_images_with_gpt(client, base64_images, prompt_type="summary", model="gpt-4o-mini", max_tokens=4000, image_detail="high", process_separately=False):
    
    if not client:
        st.error("OpenAI 클라이언트가 초기화되지 않았습니다.")
        return None
        
    if not base64_images:
        st.error("처리할 이미지가 없습니다.")
        return None
    
    # 이미지 크기 검증
    total_size = 0
    for i, img in enumerate(base64_images):
        img_size = len(img) * 3 / 4  # base64 디코딩 후 크기 추정
        total_size += img_size
        if img_size > 20 * 1024 * 1024:  # 20MB 제한
            st.error(f"이미지 {i+1}이 너무 큽니다 ({img_size/1024/1024:.1f}MB). 20MB 이하로 줄여주세요.")
            return None
    
    if total_size > 100 * 1024 * 1024:  # 전체 100MB 제한
        st.warning(f"전체 이미지 크기가 큽니다 ({total_size/1024/1024:.1f}MB). 처리 시간이 오래 걸릴 수 있습니다.")
    
    # 개별 처리 모드
    if process_separately and len(base64_images) > 1:
        st.write(f"🔍 {len(base64_images)}개 이미지를 개별적으로 처리합니다...")
        results = []
        previous_context = ""
        failed_pages = []
        
        for i, base64_img in enumerate(base64_images):
            st.write(f"📄 페이지 {i+1}/{len(base64_images)} 처리 중...")
            
            # 다음 페이지 시작 부분 미리보기
            next_page_start = ""
            if i < len(base64_images) - 1:
                try:
                    next_preview = analyze_single_image_with_context(
                        client, base64_images[i+1], prompt_type, model, 1000, "low", 
                        i+2, len(base64_images), "", ""
                    )
                    if next_preview:
                        next_start, _ = extract_overlap_context(next_preview, 300)
                        next_page_start = next_start
                except Exception as e:
                    st.warning(f"다음 페이지 미리보기 실패: {e}")
            
            # 문맥을 포함한 개별 처리
            result = analyze_single_image_with_context(
                client, base64_img, prompt_type, model, max_tokens, image_detail, 
                i+1, len(base64_images), previous_context, next_page_start
            )
            
            if result:
                results.append(result)
                previous_context = extract_context_for_next_page(result, 800)
                st.success(f"✅ 페이지 {i+1} 완료")
            else:
                failed_pages.append(i+1)
                results.append(f"❌ [페이지 {i+1} 처리 실패]")
                previous_context = ""
                st.error(f"❌ 페이지 {i+1} 처리 실패")
        
        if failed_pages:
            st.error(f"실패한 페이지: {', '.join(map(str, failed_pages))}")
        
        # 결과를 자연스럽게 연결
        final_result = "\n\n".join(results)
        
        # 전체 문서 맥락에서 최종 정리 (실패한 페이지가 적은 경우에만)
        if len(base64_images) > 2 and len(final_result) > 3000 and len(failed_pages) <= len(base64_images) * 0.3:
            st.write("🔧 전체 문서 흐름 개선 중...")
            try:
                polish_prompt = f"""다음은 페이지별로 개별 처리된 {'요약' if prompt_type == 'summary' else '번역'} 결과입니다. 
전체 문서의 맥락을 고려하여 다음을 개선해주세요:

1. 페이지 간 연결이 부자연스러운 부분만 매끄럽게 연결
2. 명백히 중복되는 내용이 있다면 정리
3. 기존 내용의 의미나 분량은 변경하지 말고, 연결부분만 개선
4. 하나의 연속된 문서로 자연스럽게 만들어주세요
- 참고문헌은 원문 그대로 유지
- 머리말/꼬리말 제거: 반복되는 책 제목, 장 제목, 페이지 번호 등은 정리

**원본 결과:**
{final_result}

**중요**: 내용을 줄이거나 생략하지 말고, 단지 페이지 간 연결만 자연스럽게 만들어서 하나의 매끄러운 문서로 만들어주세요."""

                polish_response = client.chat.completions.create(
                    model=model,
                    messages=[{"role": "user", "content": polish_prompt}],
                    max_completion_tokens=max_tokens * 2,
                    timeout=240 if model.startswith("gpt-5") else 120
                )
                
                if polish_response.choices and polish_response.choices[0].message.content:
                    polished_result = polish_response.choices[0].message.content.strip()
                    if len(polished_result) > len(final_result) * 0.8:
                        final_result = polished_result
                        st.success("✨ 최종 연결 개선 완료")
                    else:
                        st.warning("개선 결과가 너무 짧아서 원본 사용")
                
            except Exception as e:
                st.warning(f"최종 정리 중 오류 발생 (원본 결과 사용): {e}")
        
        return final_result
    
    # 기존 일괄 처리 모드
    if prompt_type == "summary":
        system_prompt = """당신은 전문 문서 요약 전문가입니다. 제공된 이미지의 내용을 분석하여 다음과 같이 요약해주세요:

📋 **요약 규칙:**
- 개조식으로 요약 (~음, ~했음 어조 사용)
- 핵심 내용과 주요 포인트 중심으로 정리
- 마크다운 형식으로 구조화
- 표나 그래프가 있다면 주요 데이터와 수치를 포함
- 이미지의 전체적인 맥락과 문서 구조 고려
- 중요한 결론이나 시사점 강조

문서의 모든 중요한 정보를 빠뜨리지 말고 체계적으로 정리해주세요."""

    else:  # translation
        system_prompt = """당신은 고급 전문 번역가입니다. 제공된 이미지의 내용을 정확하고 자연스럽게 한글로 번역해주세요.

🌐 **번역 규칙:**
1. **정확성**: 원문의 의미와 뉘앙스를 정확히 보존
2. **전문성**: 학술적/기술적 용어는 적절한 한국어 전문 용어 사용
3. **자연스러움**: 한국어 문체와 어순에 맞게 자연스럽게 번역
4. **구조 보존**: 원문의 문단 구조와 강조점 유지
5. **표와 데이터**: 표는 마크다운 표 형식으로 번역, 수치와 데이터는 정확히 보존
6. **그림과 다이어그램**: 그림 내 텍스트는 모두 번역, 그림 설명과 캡션도 번역
7. **제목과 소제목**: 적절한 한국어 제목 형식으로 번역

**중요**: 단순한 단어 치환이 아닌, 의미와 맥락을 고려한 고품질 전문 번역을 수행해주세요."""

    st.write(f"🔍 {len(base64_images)}개 이미지 일괄 처리 중...")
    
    messages = [
        {
            "role": "user",
            "content": [
                {
                    "type": "text",
                    "text": f"{system_prompt}\n\n위 규칙에 따라 다음 이미지들의 내용을 정확하고 상세하게 분석해주세요."
                }
            ] + [
                {
                    "type": "image_url",
                    "image_url": {
                        "url": f"data:image/png;base64,{base64_img}",
                        "detail": image_detail
                    }
                } for base64_img in base64_images
            ]
        }
    ]

    try:
        # GPT-5 모델 사용 시 더 높은 타임아웃 설정
        timeout_seconds = 180 if model.startswith("gpt-5") else 120
        
        response = client.chat.completions.create(
            model=model,
            messages=messages,
            max_completion_tokens=max_tokens,
            timeout=timeout_seconds
        )
        
        if response.choices and len(response.choices) > 0:
            choice = response.choices[0]
            
            if hasattr(choice, 'message') and choice.message and choice.message.content:
                content = choice.message.content.strip()
                
                if choice.finish_reason == 'length':
                    st.warning("⚠️ 응답이 토큰 제한으로 잘렸습니다. '페이지별 개별 처리' 옵션을 사용하거나 최대 토큰 수를 늘려보세요.")
                
                if content:
                    return content
                else:
                    st.error("응답 내용이 비어있습니다.")
                    return None
            else:
                st.error("API 응답 형식이 올바르지 않습니다.")
                return None
        else:
            st.error("API에서 응답을 받지 못했습니다.")
            return None
            
    except Exception as e:
        error_msg = str(e)
        if "timeout" in error_msg.lower():
            st.error("API 요청 시간 초과. 네트워크 연결을 확인하거나 이미지 수를 줄여보세요.")
        elif "rate limit" in error_msg.lower():
            st.error("API 사용량 한도 초과. 잠시 후 다시 시도해주세요.")
        else:
            st.error(f"API 오류: {error_msg}")
        return None

# QMD 파일 저장 (Quarto Markdown) - 에러 처리 개선
def save_to_qmd(filename, **sections):
    """Quarto Markdown 파일로 저장"""
    try:
        content_lines = []
        
        # YAML 헤더 추가
        content_lines.append("---")
        content_lines.append("title: \"PDF 분석 결과\"")
        content_lines.append("format: html")
        content_lines.append("editor: visual")
        content_lines.append("---")
        content_lines.append("")
        
        for title, content in sections.items():
            # 섹션 제목 추가 (레벨 1 헤딩)
            content_lines.append(f"# {title}")
            content_lines.append("")
            
            if content and content.strip():
                # 내용을 그대로 추가 (마크다운 형식 유지)
                content_lines.append(content)
            else:
                content_lines.append(f"*{title} 내용이 없습니다.*")
            
            content_lines.append("")
            content_lines.append("---")
            content_lines.append("")
        
        # 마지막 구분선 제거
        if content_lines and content_lines[-2] == "---":
            content_lines = content_lines[:-2]
        
        with open(filename, 'w', encoding='utf-8') as f:
            f.write('\n'.join(content_lines))
        return True
    except Exception as e:
        st.error(f"QMD 파일 저장 오류: {e}")
        return False

def save_to_word(filename, **sections):
    try:
        doc = Document()
        
        for title, content in sections.items():
            # 제목 추가
            heading = doc.add_heading(title, level=1)
            
            # 내용 확인 및 추가
            if content and content.strip():
                # 줄바꿈으로 분리하되 빈 줄도 유지
                lines = content.split('\n')
                
                for line in lines:
                    line_stripped = line.strip()
                    
                    # 빈 줄 처리
                    if not line_stripped:
                        doc.add_paragraph("")
                        continue
                    
                    # 마크다운 헤딩 처리
                    if line_stripped.startswith('###'):
                        doc.add_heading(line_stripped.replace('###', '').strip(), level=3)
                    elif line_stripped.startswith('##'):
                        doc.add_heading(line_stripped.replace('##', '').strip(), level=2)
                    elif line_stripped.startswith('#'):
                        doc.add_heading(line_stripped.replace('#', '').strip(), level=2)
                    else:
                        # 일반 텍스트 또는 표 처리
                        paragraph = doc.add_paragraph()
                        paragraph.add_run(line_stripped)
            else:
                # 내용이 없는 경우 알림 추가
                doc.add_paragraph(f"[{title}] 내용이 없습니다.")
            
            # 섹션 간 페이지 브레이크 (마지막 섹션 제외)
            sections_list = list(sections.items())
            if (title, content) != sections_list[-1]:
                doc.add_page_break()
        
        doc.save(filename)
        return True
    except Exception as e:
        st.error(f"Word 파일 저장 오류: {e}")
        return False

# PDF 파일 저장 (개선된 버전) - 에러 처리 강화
def save_to_pdf(filename, **sections):
    try:
        # 한글 폰트 등록 시도
        font_name = 'Helvetica'  # 기본 폰트
        try:
            # Windows 한글 폰트 시도
            font_paths = [
                'C:/Windows/Fonts/malgun.ttf',
                'C:/Windows/Fonts/batang.ttf',
                'C:/Windows/Fonts/gulim.ttf'
            ]
            for font_path in font_paths:
                if os.path.exists(font_path):
                    pdfmetrics.registerFont(TTFont('Korean', font_path))
                    font_name = 'Korean'
                    break
        except Exception as e:
            st.warning(f"한글 폰트 로딩 실패, 기본 폰트 사용: {e}")
        
        # 스타일 정의
        styles = getSampleStyleSheet()
        
        korean_normal = ParagraphStyle(
            'KoreanNormal',
            parent=styles['Normal'],
            fontName=font_name,
            fontSize=11,
            leading=16,
            leftIndent=0,
            rightIndent=0,
            spaceAfter=6
        )
        
        korean_heading = ParagraphStyle(
            'KoreanHeading',
            parent=styles['Heading1'],
            fontName=font_name,
            fontSize=16,
            leading=20,
            spaceBefore=12,
            spaceAfter=12,
            leftIndent=0
        )
        
        story = []
        
        # 내용 처리
        for title, content in sections.items():
            # 제목 추가
            story.append(Paragraph(f"<b>{title}</b>", korean_heading))
            story.append(Spacer(1, 0.2*inch))
            
            # 내용 처리
            if content and content.strip():
                # 줄 단위로 처리
                lines = content.split('\n')
                for line in lines:
                    line_clean = line.strip()
                    
                    if line_clean:
                        # HTML 이스케이프 처리
                        line_escaped = (line_clean
                                       .replace('&', '&amp;')
                                       .replace('<', '&lt;')
                                       .replace('>', '&gt;')
                                       .replace('"', '&quot;'))
                        
                        try:
                            # 마크다운 헤딩 처리
                            if line_escaped.startswith('###'):
                                heading_text = line_escaped.replace('###', '').strip()
                                story.append(Paragraph(f"<b>{heading_text}</b>", korean_normal))
                            elif line_escaped.startswith('##'):
                                heading_text = line_escaped.replace('##', '').strip()
                                story.append(Paragraph(f"<b>{heading_text}</b>", korean_normal))
                            elif line_escaped.startswith('#'):
                                heading_text = line_escaped.replace('#', '').strip()
                                story.append(Paragraph(f"<b>{heading_text}</b>", korean_normal))
                            else:
                                # 일반 텍스트
                                story.append(Paragraph(line_escaped, korean_normal))
                            
                            story.append(Spacer(1, 3))
                            
                        except Exception as e:
                            # 폰트 문제 시 기본 스타일 사용
                            try:
                                story.append(Paragraph(line_escaped, styles['Normal']))
                                story.append(Spacer(1, 3))
                            except:
                                # 그래도 안 되면 건너뛰기
                                continue
            
            # 섹션 간 페이지 브레이크
            sections_list = list(sections.items())
            if (title, content) != sections_list[-1]:
                story.append(PageBreak())
        
        doc = SimpleDocTemplate(filename, pagesize=A4)
        doc.build(story)
        return True
    except Exception as e:
        st.warning(f"PDF 생성 중 오류 발생: {e}. Word 파일을 사용해주세요.")
        return False

# 세션 상태 초기화 함수 추가
def initialize_session_state():
    """세션 상태 초기화"""
    defaults = {
        "images": [],
        "base64_images": [],
        "page_number": 1,
        "start_page": 1,
        "end_page": 1,
        "analysis_results": {},
        "last_analysis_done": False,
        "include_images": False,
        "temp_files": [],  # 임시 파일 추적용
        "api_key_validated": False,
        "client": None
    }
    
    for key, default_value in defaults.items():
        if key not in st.session_state:
            st.session_state[key] = default_value

# 메인 앱
def main():
    st.set_page_config(layout="wide", page_title="PDF Vision 번역/요약 프로그램")
    
    # 세션 상태 초기화
    initialize_session_state()

    with st.sidebar:
        st.title("📄 PDF Vision 번역/요약")
        st.markdown("*GPT Vision으로 이미지 직접 분석*")
        
        # API 키 입력 및 검증
        openai_api_key = st.text_input("OpenAI API Key", type="password", 
                                      help="OpenAI API 키를 입력하세요")
        st.write("[OpenAI API Key 받기](https://platform.openai.com/account/api-keys)")
        
        # API 키 검증
        if openai_api_key and (not st.session_state.api_key_validated or 
                              st.session_state.get('last_api_key', '') != openai_api_key):
            with st.spinner("API 키 확인 중..."):
                client, is_valid, error_msg = validate_openai_api_key(openai_api_key)
                if is_valid:
                    st.session_state.client = client
                    st.session_state.api_key_validated = True
                    st.session_state.last_api_key = openai_api_key
                    st.success("✅ API 키 확인됨")
                else:
                    st.session_state.client = None
                    st.session_state.api_key_validated = False
                    st.error(f"❌ {error_msg}")

        # 모델 선택
        model_option = st.selectbox(
            "모델 선택",
            ["gpt-4o-mini", "gpt-4o", "gpt-4o-2024-11-20", "gpt-5-mini", "gpt-5"],
            index=0,
            help="GPT-5 계열은 최고 품질이지만 더 많은 토큰을 사용합니다. GPT-4o는 균형잡힌 성능을 제공합니다."
        )

        pdf_file = st.file_uploader("PDF 파일 업로드", type=["pdf"])
        mode = st.radio("페이지 선택", ["단일 페이지", "페이지 범위", "전체 문서"])

        client = st.session_state.client

        # PDF 파일 처리
        if pdf_file:
            # 파일이 변경되었는지 확인
            current_file_hash = hash(pdf_file.getvalue())
            if st.session_state.get('last_file_hash') != current_file_hash:
                # 이전 임시 파일들 정리
                cleanup_temp_files(st.session_state.temp_files)
                st.session_state.temp_files = []
                
                pdf_data = pdf_file.read()
                with st.spinner("PDF를 이미지로 변환 중..."):
                    st.session_state.images, st.session_state.base64_images = convert_pdf_to_base64_images(pdf_data)
                
                # 이미지 경로들을 임시 파일 목록에 추가
                st.session_state.temp_files.extend(st.session_state.images)
                st.session_state.last_file_hash = current_file_hash
                
                # 분석 결과 초기화
                st.session_state.analysis_results = {}
                st.session_state.last_analysis_done = False
            
            total_pages = len(st.session_state.images)
            if total_pages > 0:
                st.success(f"총 {total_pages}페이지 변환 완료")

                if mode == "단일 페이지":
                    st.session_state.page_number = st.number_input(
                        "페이지 번호", 1, total_pages, st.session_state.page_number
                    )
                elif mode == "페이지 범위":
                    st.session_state.start_page = st.number_input(
                        "시작 페이지", 1, total_pages, st.session_state.start_page
                    )
                    st.session_state.end_page = st.number_input(
                        "끝 페이지", st.session_state.start_page, total_pages,
                        min(st.session_state.end_page, total_pages)
                    )
            else:
                st.error("PDF 변환에 실패했습니다.")

    # 메인 콘텐츠
    if pdf_file and st.session_state.images:
        left, right = st.columns([1, 1])
        
        with left:
            st.subheader("📖 문서 미리보기")
            try:
                if mode == "단일 페이지":
                    if 0 <= st.session_state.page_number-1 < len(st.session_state.images):
                        st.image(st.session_state.images[st.session_state.page_number-1], 
                                caption=f"페이지 {st.session_state.page_number}")
                elif mode == "페이지 범위":
                    for i in range(st.session_state.start_page-1, min(st.session_state.end_page, len(st.session_state.images))):
                        if i < len(st.session_state.images):
                            st.image(st.session_state.images[i], caption=f"페이지 {i+1}")
                else:
                    # 전체 문서의 경우 처음 5페이지만 미리보기
                    preview_count = min(5, len(st.session_state.images))
                    for i in range(preview_count):
                        st.image(st.session_state.images[i], caption=f"페이지 {i+1}")
                    if len(st.session_state.images) > 5:
                        st.info(f"미리보기는 처음 5페이지만 표시됩니다. (전체 {len(st.session_state.images)}페이지)")
            except Exception as e:
                st.error(f"이미지 표시 오류: {e}")

        with right:
            st.subheader("🤖 AI 분석 및 처리")
            
            if not st.session_state.api_key_validated:
                st.warning("먼저 유효한 OpenAI API 키를 입력해주세요.")
            else:
                # 처리 옵션
                col1, col2 = st.columns(2)
                with col1:
                    do_summary = st.checkbox("📋 요약하기", value=True)
                with col2:
                    do_translation = st.checkbox("🌐 번역하기", value=True)
                
                # 추가 옵션
                with st.expander("🔧 고급 설정"):
                    include_images = st.checkbox("📸 결과에 원본 이미지 포함", value=False, 
                                               help="번역/요약 결과와 함께 원본 이미지도 표시합니다")
                    
                    # 모델별 권장 토큰 수 설정
                    if model_option in ["gpt-5", "gpt-5-mini"]:
                        default_tokens = 12000
                        max_tokens_limit = 32000
                        token_help = "GPT-5 계열은 더 많은 토큰이 필요하며, 더 긴 응답을 생성할 수 있습니다."
                    else:
                        default_tokens = 8000
                        max_tokens_limit = 16000
                        token_help = "GPT-4o 계열의 권장 토큰 수입니다."
                    
                    max_tokens = st.slider("최대 토큰 수", 2000, max_tokens_limit, default_tokens, 
                                         help=token_help)
                    
                    # 이미지 처리 방식 선택
                    image_detail = st.selectbox(
                        "이미지 해상도",
                        ["high", "low"],
                        index=0,
                        help="high: 고화질 분석 (토큰 많이 사용), low: 저화질 분석 (토큰 절약)"
                    )
                    
                    # 페이지 분할 처리 옵션
                    process_separately = st.checkbox(
                        "페이지별 개별 처리 (문맥 연결)", 
                        value=True,
                        help="여러 페이지를 개별적으로 처리하되, 페이지 간 문맥을 연결하여 자연스럽게 처리 (권장)"
                    )
                    
                    # GPT-5 사용 시 추가 안내
                    if model_option in ["gpt-5", "gpt-5-mini"]:
                        st.info("💡 GPT-5 모델 사용 시 더 정확하고 자연스러운 번역/요약이 가능하지만, 처리 시간과 비용이 증가할 수 있습니다.")
                    
                    # include_images 값을 세션에 저장
                    st.session_state.include_images = include_images

                if not (do_summary or do_translation):
                    st.warning("요약 또는 번역 중 하나는 선택해주세요.")
                
                # 분석 시작 버튼
                analysis_disabled = not (do_summary or do_translation) or not st.session_state.api_key_validated
                if st.button("🚀 AI 분석 시작", disabled=analysis_disabled):
                    if not st.session_state.base64_images:
                        st.error("처리할 이미지가 없습니다.")
                    else:
                        # 선택된 페이지의 이미지만 추출
                        try:
                            if mode == "단일 페이지":
                                if 0 <= st.session_state.page_number-1 < len(st.session_state.base64_images):
                                    selected_images = [st.session_state.base64_images[st.session_state.page_number-1]]
                                else:
                                    st.error("선택한 페이지가 범위를 벗어났습니다.")
                                    selected_images = []
                            elif mode == "페이지 범위":
                                start_idx = max(0, st.session_state.start_page-1)
                                end_idx = min(len(st.session_state.base64_images), st.session_state.end_page)
                                selected_images = st.session_state.base64_images[start_idx:end_idx]
                            else:
                                selected_images = st.session_state.base64_images

                            if not selected_images:
                                st.error("선택된 페이지가 없습니다.")
                            else:
                                results = {}
                                
                                # 요약 처리
                                if do_summary:
                                    with st.spinner("📋 문서 요약 중..."):
                                        summary = analyze_images_with_gpt(
                                            client=client, 
                                            base64_images=selected_images, 
                                            prompt_type="summary", 
                                            model=model_option, 
                                            max_tokens=max_tokens, 
                                            image_detail=image_detail, 
                                            process_separately=process_separately
                                        )
                                        if summary:
                                            results["요약 결과"] = summary

                                # 번역 처리
                                if do_translation:
                                    with st.spinner("🌐 문서 번역 중..."):
                                        translation = analyze_images_with_gpt(
                                            client=client,
                                            base64_images=selected_images, 
                                            prompt_type="translation", 
                                            model=model_option,
                                            max_tokens=max_tokens, 
                                            image_detail=image_detail, 
                                            process_separately=process_separately
                                        )
                                        if translation:
                                            results["번역 결과"] = translation

                                # 결과를 세션 상태에 저장
                                if results:
                                    st.session_state.analysis_results = results
                                    st.session_state.last_analysis_done = True
                                    st.success("✅ 분석 완료!")
                                else:
                                    st.error("분석 결과가 없습니다. 다시 시도해주세요.")
                        
                        except Exception as e:
                            st.error(f"분석 중 오류 발생: {e}")
    
    # 분석 결과 표시 (세션 상태에서 가져오기)
    if st.session_state.last_analysis_done and st.session_state.analysis_results:
        results = st.session_state.analysis_results
        
        st.markdown("---")
        
        # 요약 결과 표시
        if "요약 결과" in results:
            st.subheader("📋 요약 결과")
            st.markdown(results["요약 결과"])
            
            # 원본 이미지 포함 옵션 (요약용)
            if st.session_state.get('include_images', False):
                st.subheader("📸 원본 이미지 (요약)")
                try:
                    if mode == "단일 페이지":
                        if 0 <= st.session_state.page_number-1 < len(st.session_state.images):
                            st.image(st.session_state.images[st.session_state.page_number-1], 
                                    caption=f"원본 페이지 {st.session_state.page_number}")
                    elif mode == "페이지 범위":
                        for i in range(st.session_state.start_page-1, min(st.session_state.end_page, len(st.session_state.images))):
                            if i < len(st.session_state.images):
                                st.image(st.session_state.images[i], caption=f"원본 페이지 {i+1}")
                    else:
                        for i, img in enumerate(st.session_state.images):
                            st.image(img, caption=f"원본 페이지 {i+1}")
                except Exception as e:
                    st.error(f"원본 이미지 표시 오류: {e}")

        # 번역 결과 표시
        if "번역 결과" in results:
            st.subheader("🌐 번역 결과")
            st.markdown(results["번역 결과"])
            
            # 원본 이미지 포함 옵션 (번역용 - 요약이 없는 경우에만)
            if st.session_state.get('include_images', False) and "요약 결과" not in results:
                st.subheader("📸 원본 이미지 (번역)")
                try:
                    if mode == "단일 페이지":
                        if 0 <= st.session_state.page_number-1 < len(st.session_state.images):
                            st.image(st.session_state.images[st.session_state.page_number-1], 
                                    caption=f"원본 페이지 {st.session_state.page_number}")
                    elif mode == "페이지 범위":
                        for i in range(st.session_state.start_page-1, min(st.session_state.end_page, len(st.session_state.images))):
                            if i < len(st.session_state.images):
                                st.image(st.session_state.images[i], caption=f"원본 페이지 {i+1}")
                    else:
                        for i, img in enumerate(st.session_state.images):
                            st.image(img, caption=f"원본 페이지 {i+1}")
                except Exception as e:
                    st.error(f"원본 이미지 표시 오류: {e}")

        # 파일 다운로드 섹션
        st.subheader("💾 결과 다운로드")
        
        temp_files_for_cleanup = []
        
        try:
            # Word 파일 생성
            with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp_docx:
                if save_to_word(tmp_docx.name, **results):
                    tmp_docx_path = tmp_docx.name
                    temp_files_for_cleanup.append(tmp_docx_path)
                else:
                    tmp_docx_path = None

            # PDF 파일 생성
            with tempfile.NamedTemporaryFile(delete=False, suffix=".pdf") as tmp_pdf:
                if save_to_pdf(tmp_pdf.name, **results):
                    tmp_pdf_path = tmp_pdf.name
                    temp_files_for_cleanup.append(tmp_pdf_path)
                else:
                    tmp_pdf_path = None
                
            # QMD 파일 생성
            with tempfile.NamedTemporaryFile(delete=False, suffix=".qmd") as tmp_qmd:
                if save_to_qmd(tmp_qmd.name, **results):
                    tmp_qmd_path = tmp_qmd.name
                    temp_files_for_cleanup.append(tmp_qmd_path)
                else:
                    tmp_qmd_path = None

            col1, col2, col3, col4 = st.columns([1, 1, 1, 1])
            
            with col1:
                if tmp_docx_path:
                    try:
                        with open(tmp_docx_path, "rb") as f:
                            docx_data = f.read()
                            st.download_button(
                                "📄 Word 다운로드", 
                                docx_data, 
                                file_name="pdf_analysis_result.docx",
                                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                            )
                    except Exception as e:
                        st.error(f"Word 파일 읽기 오류: {e}")
                else:
                    st.error("Word 파일 생성 실패")
                    
            with col2:
                if tmp_pdf_path:
                    try:
                        with open(tmp_pdf_path, "rb") as f:
                            pdf_data = f.read()
                            st.download_button(
                                "📕 PDF 다운로드", 
                                pdf_data, 
                                file_name="pdf_analysis_result.pdf",
                                mime="application/pdf"
                            )
                    except Exception as e:
                        st.error(f"PDF 파일 읽기 오류: {e}")
                else:
                    st.warning("PDF 파일 생성 실패")
                    
            with col3:
                if tmp_qmd_path:
                    try:
                        with open(tmp_qmd_path, "rb") as f:
                            qmd_data = f.read()
                            st.download_button(
                                "📝 QMD 다운로드", 
                                qmd_data, 
                                file_name="pdf_analysis_result.qmd",
                                mime="text/plain"
                            )
                    except Exception as e:
                        st.error(f"QMD 파일 읽기 오류: {e}")
                else:
                    st.error("QMD 파일 생성 실패")
                    
            with col4:
                if st.button("🗑️ 결과 지우기"):
                    st.session_state.analysis_results = {}
                    st.session_state.last_analysis_done = False
                    st.rerun()

        except Exception as e:
            st.error(f"파일 생성 중 오류 발생: {e}")
        
        finally:
            # 임시 파일 정리
            cleanup_temp_files(temp_files_for_cleanup)

    # 사용법 안내
    if not pdf_file:
        st.markdown("""
        ## 🎯 사용법
        
        1. **OpenAI API Key 입력**: 사이드바에서 API 키를 입력하세요
        2. **PDF 업로드**: 번역/요약하고 싶은 PDF 파일을 업로드하세요
        3. **페이지 선택**: 전체 문서 또는 특정 페이지를 선택하세요
        4. **처리 옵션 선택**: 요약, 번역 또는 둘 다 선택하세요
        5. **AI 분석 시작**: 버튼을 눌러 GPT Vision으로 분석하세요
        
        ## ✨ 특징
        
        - **OCR 불필요**: GPT Vision이 이미지를 직접 읽어 더 정확함
        - **시각적 요소 분석**: 표, 그래프, 다이어그램도 함께 분석
        - **컨텍스트 이해**: 문서 구조와 레이아웃을 고려한 분석
        - **고품질 번역/요약**: GPT-5의 최고급 언어 능력 또는 GPT-4o의 균형잡힌 성능 활용
        - **문맥 연결**: 페이지 간 이어지는 내용을 자연스럽게 연결 처리
        - **다양한 모델**: GPT-4o부터 최신 GPT-5까지 다양한 모델 선택 가능
        - **안정성 향상**: 에러 처리 및 메모리 관리 개선
        """)

# 앱 종료 시 정리 작업
@st.cache_resource
def cleanup_on_exit():
    """앱 종료 시 임시 파일 정리"""
    def cleanup():
        if 'temp_files' in st.session_state:
            cleanup_temp_files(st.session_state.temp_files)
    
    import atexit
    atexit.register(cleanup)
    return True

if __name__ == "__main__":
    cleanup_on_exit()
    main()
